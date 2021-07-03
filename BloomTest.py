import os
import time
import docx

from bloomfilter import BloomFilter

# Каталог из которого будем брать файлы
directory = 'D:/iliya/Documents/Политех/3 курс/2 семестр/Нечёткая логика/4лаба'

#Слова
FindWords = ['Министерство', 'Опоссум', 'Карпов', 'система', 'диаграмма',
             'вариант', 'Яблоко', 'Интеграл', 'фамилия', 'лицо', 'рука', 'Paradigm', 'ГОСТ', 'освещение']

# Получаем список файлов в переменную files
files = os.listdir(directory)
bf = BloomFilter(100000, 0.01)
t0 = time.time_ns()
# По фильтру Блума
for file in files:
    path = directory + '/' + file
    doc = docx.Document(path)
    for paragraph in doc.paragraphs:
        for word in paragraph.text.split():
            bf.add(word)
tblum = time.time_ns()-t0
print("ФИЛЬТР БЛУМА")
for word in FindWords:
    t0 = time.time_ns()
    if bf.check(word):
        print("YES")
    else:
        print("NO")
    print("Время работы: ", time.time_ns() - t0 + tblum)
print('\n')
print("Перебор")
#По классика
arr = []
t0 = time.time_ns()
for file in files:
    path = directory + '/' + file
    doc = docx.Document(path)
    for paragraph in doc.paragraphs:
        for word in paragraph.text.split():
            arr.append(word)
tperebor = time.time_ns() - t0
for word in FindWords:
    t0 = time.time_ns()
    if word in arr:
        print("YES")
    else:
        print("NO")
    print("Время работы: ", time.time_ns() - t0 + tperebor)


